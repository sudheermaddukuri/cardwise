"""
Generates CardWise_RAG_Project.docx in the project root.
Run: uv run python scripts/generate_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path

OUT = Path(__file__).parent.parent / "CardWise_RAG_Project.docx"

# ── helpers ───────────────────────────────────────────────────────────────────

def set_font(run, size=11, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        run.font.size = Pt(11)
    return p

def bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix + " ")
        set_font(r, bold=True)
        r2 = p.add_run(text)
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.size = Pt(10)
            para.paragraph_format.space_after = Pt(0)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "2563EB")
        shd.set(qn("w:color"), "FFFFFF")
        shd.set(qn("w:val"), "clear")
        tcPr.append(shd)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = str(val)
            for para in row_cells[c_idx].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10)
                para.paragraph_format.space_after = Pt(0)
            if r_idx % 2 == 0:
                tc = row_cells[c_idx]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:fill"), "EFF6FF")
                shd.set(qn("w:val"), "clear")
                tcPr.append(shd)

    if col_widths:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[i])

    doc.add_paragraph()


def code_block(doc, lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        run = p.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(30, 41, 59)
    doc.add_paragraph()


# ── document ──────────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Default style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# ── Cover ─────────────────────────────────────────────────────────────────────

doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("CardWise — AI-Powered Credit Card Advisor")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(37, 99, 235)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = subtitle.add_run("Building a RAG System from Scratch\nProject Writeup — June 2026")
r2.font.size = Pt(12)
r2.font.color.rgb = RGBColor(100, 116, 139)

doc.add_paragraph()
doc.add_paragraph()

# ── 1. Problem Statement ──────────────────────────────────────────────────────

heading(doc, "1. Problem Statement")

body(doc,
    "Most people own multiple credit cards but have no easy way to know whether "
    "they are actually getting good value from them. The reward structures are complex "
    "— rotating quarterly categories, point currencies with varying redemption rates, "
    "streaming bonuses tied to specific merchants — and understanding whether your "
    "spending patterns align with your card's best earn rates requires time and effort "
    "most people don't have."
)

body(doc,
    "The goal of this project was to build an AI assistant that can answer natural "
    "language questions like:"
)

bullet(doc, '"Am I getting good value from my dining spend?"')
bullet(doc, '"Which card should I use for gas this month?"')
bullet(doc, '"Should I book a concert ticket to maximise rewards?"')
bullet(doc, '"What am I leaving on the table with my current card usage?"')

body(doc,
    "The assistant should answer using the user's actual spending history and exact card "
    "terms — not generic advice. This is the core challenge: grounding AI answers in "
    "real, user-specific data."
)

# ── 2. Why This Is Hard Without RAG ──────────────────────────────────────────

heading(doc, "2. Why This Is Hard Without AI")

body(doc,
    "A plain GPT-4 call cannot solve this because:"
)

bullet(doc, "it has no knowledge of a user's actual spend history")
bullet(doc, "it may hallucinate card terms (wrong rates, wrong categories)")
bullet(doc, "it cannot compare what a user earned vs what they could have earned")

body(doc,
    "What we need is a way to inject the right facts into the model at query time — "
    "pulling from the user's own data and the exact card terms, without the model "
    "inventing anything. This is exactly what RAG (Retrieval-Augmented Generation) solves."
)

# ── 3. What is RAG ───────────────────────────────────────────────────────────

heading(doc, "3. What is RAG?")

body(doc,
    "RAG stands for Retrieval-Augmented Generation. The idea is simple: before asking "
    "the LLM a question, you first retrieve the relevant facts from a database and "
    "attach them to the prompt. The LLM then answers using those facts as context rather "
    "than relying on its training data."
)

body(doc, "The three-step pattern:")

bullet(doc, "the user's spending data and card reward terms", bold_prefix="Store —")
bullet(doc, "find the chunks most relevant to the user's question", bold_prefix="Retrieve —")
bullet(doc, "pass retrieved context + question to the LLM; get a grounded answer", bold_prefix="Generate —")

body(doc,
    "The result is an assistant that can say '\"You earned 2% cashback on dining with "
    "Citi, but your Chase card earns 3x points — that is a better rate\"' because it "
    "retrieved both facts from your data, not because it guessed."
)

# ── 4. Solution Architecture ─────────────────────────────────────────────────

heading(doc, "4. Solution Architecture")

body(doc,
    "The system is built in four sequential steps, each building on the last."
)

add_table(doc,
    headers=["Step", "What it does", "Output"],
    rows=[
        ["1  Data Generation",
         "Simulate 6 months of realistic credit card statements for 3 user personas. "
         "Each transaction is tied to a card, category, merchant, and the exact reward earned.",
         "18 JSON statement files"],
        ["2  Indexing",
         "Convert statements and card reward terms into text chunks. Embed each chunk "
         "with OpenAI text-embedding-ada-002. Store vectors + metadata in ChromaDB.",
         "118 vectors in ChromaDB"],
        ["3  Retrieval + Chain",
         "Embed the user's query. Find the closest matching chunks (filtered by user). "
         "Assemble a prompt with context and send to GPT-4o.",
         "Grounded natural language answer"],
        ["4  Chat UI",
         "Streamlit chat interface with persona selector, starter prompts, and "
         "multi-turn conversation history.",
         "Interactive web app"],
    ],
    col_widths=[1.5, 3.5, 1.8]
)

# ── 5. Key Concepts ───────────────────────────────────────────────────────────

heading(doc, "5. Key Concepts Learned")

heading(doc, "5.1 Chunking", level=2)
body(doc,
    "A chunk is the unit of retrieval — the smallest piece of text that gets embedded "
    "and stored. We used two types:"
)
bullet(doc,
    "One chunk per (user × card × category × month). "
    "Example: \"Alice spent $293 on dining using Citi Double Cash in Dec 2025 "
    "and earned $5.87 cashback at 2%.\"",
    bold_prefix="Spend summaries —"
)
bullet(doc,
    "One chunk per card encoding its reward rules as prose. "
    "Example: \"American Express Gold earns 4x Membership Rewards on dining, "
    "valued at 1¢ per point.\"",
    bold_prefix="Card rules —"
)
body(doc,
    "Retrieving both types lets the LLM compare what actually happened against what "
    "the card promises — which is what enables recommendations."
)

heading(doc, "5.2 Embeddings", level=2)
body(doc,
    "An embedding converts a text string into a list of 1,536 numbers (a vector). "
    "The key property: semantically similar texts produce vectors that are close together "
    "in that 1,536-dimensional space, even if they use different words. "
    "\"Dining spend\" and \"restaurant purchases\" land near each other; "
    "\"annual fee\" lands far away. This is why semantic search finds relevant chunks "
    "without requiring keyword matches."
)

heading(doc, "5.3 Vector Database (ChromaDB)", level=2)
body(doc,
    "ChromaDB stores each vector alongside its original text and metadata tags "
    "(user_id, card_id, category, month). At query time, it finds the closest vectors "
    "to the query using cosine similarity, and metadata filters (e.g. user_id = alice) "
    "ensure one user's data never leaks into another's results."
)

heading(doc, "5.4 Prompt Assembly", level=2)
body(doc,
    "The LLM only sees what we hand it. Before calling GPT-4o, we assemble a prompt "
    "with three sections:"
)
bullet(doc, "defines the assistant's role and prohibits inventing card terms", bold_prefix="System prompt —")
bullet(doc, "the retrieved spend summaries and card rules", bold_prefix="Context block —")
bullet(doc, "the user's actual question", bold_prefix="User question —")
body(doc,
    "Without the context block, the LLM has no way to reference Alice's data or the "
    "exact card rates — it would either hallucinate or give generic advice."
)

heading(doc, "5.5 Multi-turn Conversation", level=2)
body(doc,
    "Each API call is stateless. To support follow-up questions, prior turns are "
    "passed as a message list with each new query. The conversation history handles "
    "\"you just told me X\" continuity; retrieval runs fresh on each turn to find "
    "the most relevant chunks for the current question."
)

# ── 6. Data Design ────────────────────────────────────────────────────────────

heading(doc, "6. Data Design")

body(doc, "Three personas were created to cover different card usage patterns:")

add_table(doc,
    headers=["Persona", "Cards Held", "Spending Focus"],
    rows=[
        ["Alice Chen",   "Citi Double Cash + Chase Sapphire Preferred", "Mixed — dining, groceries, travel, shopping"],
        ["Bob Martinez", "Amex Gold",                                    "Dining and groceries heavy"],
        ["Carol Kim",    "Discover it Cash Back",                       "Gas, groceries, rotating categories"],
    ],
    col_widths=[1.3, 2.8, 2.7]
)

body(doc,
    "Each persona's statements span December 2025 through May 2026 — 6 months, "
    "18 files total. Transactions are randomised with seasonal variance, merchant-level "
    "amount ranges, and weekend bias for dining. Reward computations (cashback, points, "
    "rotating category caps) are driven entirely by machine-readable rules in cards.json "
    "— no hardcoded logic per card."
)

# ── 7. Sample Interaction ─────────────────────────────────────────────────────

heading(doc, "7. Sample Interaction")

body(doc, "Query sent to the system:")

code_block(doc, [
    'User (Alice): "Am I getting good value from my dining spend?"',
])

body(doc, "What happens internally:")

bullet(doc, "query is embedded into a 1,536-dim vector")
bullet(doc, "ChromaDB returns Alice's top dining spend chunks + card rule chunks")
bullet(doc, "context + query assembled into a prompt and sent to GPT-4o")
bullet(doc, "GPT-4o returns a grounded answer citing exact rates and amounts")

body(doc, "Actual response from the system:")

code_block(doc, [
    '"You\'re currently using the Chase Sapphire Preferred for most of your',
    ' dining spend, earning 3x points at 1.25¢/pt — an effective 3.75% return.',
    ' However, the American Express Gold earns 4x Membership Rewards on dining',
    ' at 1¢/pt — a 4% return. Routing dining to Amex Gold would improve your',
    ' reward rate on that category."',
])

body(doc,
    "Every figure in that answer was retrieved from the index, not generated from "
    "the model's training. This is the core value of RAG."
)

# ── 8. Tech Stack ─────────────────────────────────────────────────────────────

heading(doc, "8. Technology Stack")

add_table(doc,
    headers=["Component", "Tool", "Why"],
    rows=[
        ["Application framework", "Streamlit",             "Rapid Python web UI — no frontend code"],
        ["Embedding model",       "text-embedding-ada-002","Industry standard, 1,536 dims, low cost"],
        ["LLM",                   "GPT-4o",               "Fast, accurate, follows system prompt rules well"],
        ["Vector database",       "ChromaDB",             "Local, persistent, simple Python API — visible internals"],
        ["Data validation",       "Pydantic",             "Typed card models, schema enforcement"],
        ["Dependency management", "uv",                   "Fast, reproducible Python environments"],
        ["Learning environment",  "Jupyter notebooks",    "Step-by-step RAG concepts with live outputs"],
    ],
    col_widths=[2.0, 2.0, 2.7]
)

# ── 9. What This Project Demonstrates ────────────────────────────────────────

heading(doc, "9. What This Project Demonstrates")

bullet(doc,
    "RAG is not a black box — every step (chunk, embed, store, retrieve, prompt, generate) "
    "is explicit, inspectable, and independently testable."
)
bullet(doc,
    "Chunking strategy is the most important design decision in a RAG system. "
    "The wrong chunk size or structure produces irrelevant retrievals regardless of how "
    "good the LLM is."
)
bullet(doc,
    "Metadata filtering is how you isolate users in a shared vector store — "
    "one collection, many users, zero data leakage."
)
bullet(doc,
    "Grounding matters. Without retrieved context, LLMs give generic answers or "
    "hallucinate specific figures. With it, every claim traces back to a real data point."
)
bullet(doc,
    "The same RAG pattern scales — swap the statements for real bank exports, "
    "swap ChromaDB for a hosted vector DB, and the architecture is production-ready."
)

doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = footer.add_run("Built with Python · ChromaDB · OpenAI · Streamlit")
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(148, 163, 184)

doc.save(OUT)
print(f"Saved: {OUT}")
